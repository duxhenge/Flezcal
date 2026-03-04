#!/usr/bin/env python3
"""Generate a nicely formatted Flezcal User Manual as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Define styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title style
title_style = doc.styles['Title']
title_style.font.name = 'Calibri'
title_style.font.size = Pt(28)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0xE8, 0x6B, 0x1A)  # Flezcal orange
title_style.paragraph_format.space_after = Pt(4)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0xE8, 0x6B, 0x1A)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(8)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(15)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(6)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
h3.paragraph_format.space_before = Pt(14)
h3.paragraph_format.space_after = Pt(4)


def add_body(text):
    """Add a body paragraph, handling **bold** inline."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def add_bullet(text, level=0):
    """Add a bullet point, handling **bold** inline."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'

    doc.add_paragraph()  # spacer


def add_divider():
    """Add a visual divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _')
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


# ============================================================
# TITLE PAGE
# ============================================================

# Spacer
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Flezcal')
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(0xE8, 0x6B, 0x1A)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Manual')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Find your favorite foods and drinks, wherever you go.')
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Calibri'

for _ in range(6):
    doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 1.0')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================

doc.add_heading('Contents', level=1)

toc_items = [
    '1. What is Flezcal?',
    '2. Web Checks vs. Community Verification',
    '3. Getting Started',
    '4. Your Picks: Choosing What You\'re Looking For',
    '5. The Map',
    '6. Adding a Spot',
    '7. Spot Details',
    '8. Rating a Spot',
    '9. The Spots Tab',
    '10. Leaderboard',
    '11. Profile',
    '12. For Venue Owners',
    '13. Tips and Tricks',
    '14. Category Reference',
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. WHAT IS FLEZCAL?
# ============================================================

doc.add_heading('1. What is Flezcal?', level=1)

add_body('Flezcal is a community-powered guide for finding specific foods and drinks you\'re passionate about. The name combines "flan" and "mezcal" \u2014 two of the original categories \u2014 but the concept extends far beyond those two items. A "Flezcal" is any specific food or drink you love and want to track down wherever you go.')

add_body('The app helps you answer one question: **"Does this place have what I\'m looking for?"**')

add_body('Rather than reviewing entire restaurants, Flezcal focuses on individual items. You might want to know which bars near you carry mezcal, which bakeries make flan, or which restaurants serve wood-fired pizza. Flezcal finds those places, checks their websites automatically, and lets the community verify and rate them.')

add_divider()

# ============================================================
# 2. WEB CHECKS VS COMMUNITY VERIFICATION
# ============================================================

doc.add_heading('2. Web Checks vs. Community Verification', level=1)

add_body('Flezcal uses two systems to help you find spots, and it\'s important to understand what each one does.')

add_body('**Web checks** are automated scans that happen when you tap on a suggested pin. The app visits the venue\'s website and searches for keywords related to your picks. This is a best-effort scan \u2014 it works well when restaurants list their menu items online, but it can miss items that aren\'t on the website, and occasionally it may match a keyword out of context.')

add_body('**Community verification** is powered by real people. When a user has physically been to a spot and confirms it serves a particular item, that carries more weight than any automated scan. Users can vote thumbs-up or thumbs-down on each category at a spot, building a trust score over time.')

add_body('Think of web checks as a helpful first guess, and community verification as the real answer. Both work together \u2014 the web check helps you decide whether a place is worth visiting, and your verification after visiting helps everyone else.')

add_divider()

# ============================================================
# 3. GETTING STARTED
# ============================================================

doc.add_heading('3. Getting Started', level=1)

doc.add_heading('Age Verification', level=2)
add_body('The first time you open Flezcal, you\'ll see an age verification screen asking "Are you 21 or older?" Because the app includes alcohol-related categories (mezcal, bourbon, scotch, etc.), you must confirm you\'re 21+ to proceed. Tap **"Yes, I\'m 21+"** to continue. This only appears once.')

doc.add_heading('Welcome Screen', level=2)
add_body('After age verification, a welcome screen introduces the app\'s key features:')
add_bullet('Discover Flezcals on an interactive map')
add_bullet('Add your favorite spots')
add_bullet('Rate your finds')
add_bullet('Climb the leaderboard as you contribute')
add_bullet('Ghost pins show unconfirmed spots \u2014 help verify them')
add_body('Tap **"Let\'s go!"** to enter the app. This screen may reappear when new features are added \u2014 you can also revisit it anytime from **Profile > What\'s New**.')

doc.add_heading('Signing In', level=2)
add_body('You can browse the map and explore without signing in, but you\'ll need an account to add spots, rate them, or verify categories. Flezcal offers two sign-in options:')
add_bullet('**Sign in with Apple** \u2014 the fastest option, uses your Apple ID')
add_bullet('**Continue with Email** \u2014 create an account with your email and a password')
add_body('You can sign in from the **Profile** tab or when you try to add a spot.')

add_divider()

# ============================================================
# 4. YOUR PICKS
# ============================================================

doc.add_heading('4. Your Picks: Choosing What You\'re Looking For', level=1)

add_body('Your picks are the heart of the app. They determine what shows up on your map, which ghost pins appear, and what the web checker searches for. You can have up to **3 picks** at a time \u2014 any combination of built-in and custom categories.')

doc.add_heading('Viewing and Managing Picks', level=2)
add_body('Go to the **My Flezcals** tab (heart icon) to see your current picks. New users start with three defaults: Mezcal, Flan, and Handmade Tortillas.')
add_body('Each pick card shows:')
add_bullet('The category emoji and name')
add_bullet('The first few search terms used to find it on websites')
add_bullet('An **edit button** (slider icon) to customize search terms')
add_bullet('A **remove button** (X icon) to remove it from your picks')
add_body('Below your active picks, empty slots show how many more you can add. Tap an empty slot or the **"Customize My Flezcals"** button to open the full category grid.')

doc.add_heading('Available Categories', level=2)
add_body('Flezcal launches with 23 curated categories across several groups:')
add_body('**The Launch Trio:** Mezcal, Flan, Handmade Tortillas')
add_body('**Drinks:** Bourbon, Fernet Branca, New England IPA, Single Malt Scotch')
add_body('**Savory:** Peameal Bacon, Wood-Fired Pizza, Paella, Oysters, Pho, Pozole, Tartare, Fugu, Bibimbap, Iberico Ham, Caviar, Pierogi, Lobster Rolls, Smashburgers')
add_body('**Sweets & Specialty:** Maple Syrup, Artisan Chocolate')
add_body('To add or remove a category, tap it in the grid. Selected categories show a checkmark. A counter at the top shows how many slots you\'ve used.')

doc.add_heading('Creating a Custom Category', level=2)
add_body('Can\'t find what you\'re looking for? Tap **"Create Your Own"** at the bottom of the category grid. Custom categories count toward your 3-pick limit, so you\'ll need to remove an existing pick if all slots are full.')
add_bullet('**Name your category** \u2014 be specific. "Pupusas" or "Empanadas" works great. Broad terms like "Italian" or "Seafood" are too generic and will be rejected.')
add_bullet('**Choose an emoji** from the provided selection.')
add_bullet('**Review the search terms** \u2014 the app auto-generates terms based on your category name. You can add, remove, or modify these.')
add_bullet('Tap **"Create & Add to My Flezcals"** to add it to your picks.')
add_body('Custom categories appear with a purple color scheme to distinguish them from built-in categories.')

doc.add_heading('Adjusting Web Search Terms', level=2)
add_body('Every category has a set of search terms that the app uses when scanning venue websites. You can customize these for any pick \u2014 built-in or custom.')
add_body('**To edit search terms:**')
add_bullet('Go to the **My Flezcals** tab')
add_bullet('Tap the **edit button** (slider icon) on any pick card')
add_bullet('The Edit Search Terms screen shows your current terms as removable chips')
add_body('**To remove a term:** Tap the X on any chip.')
add_body('**To add a term:** Type a new term in the text field and tap the + button. Think about words that restaurants actually use on their menus.')
add_body('**To reset:** Tap **"Reset to defaults"** to restore the original search terms.')
add_body('Tap **"Save Changes"** when done. Future website scans will use your updated terms.')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Tips for effective search terms:')
run.bold = True
add_bullet('Use words that appear on actual restaurant menus and websites')
add_bullet('Include common alternate spellings (e.g., "pierogi" and "pierogies")')
add_bullet('Include terms in other languages if relevant (e.g., "flan casero" for flan)')
add_bullet('Avoid overly common words that would match unrelated content')

add_divider()

# ============================================================
# 5. THE MAP
# ============================================================

doc.add_heading('5. The Map', level=1)

add_body('The **Explore** tab (map icon) is the main screen. It shows three types of markers:')

doc.add_heading('Your Location', level=3)
add_body('A standard blue dot shows where you are.')

doc.add_heading('Confirmed Spots (Solid Pins)', level=3)
add_body('These are venues that community members have added to Flezcal. They have solid, colored pins with the category emoji. Tap one to see full details, ratings, and offerings.')

doc.add_heading('Ghost Pins (Suggested Spots)', level=3)
add_body('Ghost pins are unconfirmed suggestions pulled from Apple Maps. They represent places that might have what you\'re looking for but haven\'t been verified yet.')
add_body('**Yellow ghost pins** (dashed border with "?" icon) are untested suggestions. The app hasn\'t checked their website yet.')
add_body('**Green ghost pins** (solid border with checkmark, gently pulsing) are likely matches. During a batch pre-screen, the app did a quick homepage scan and found keywords related to one of your picks. These are worth checking first.')

doc.add_heading('Using the Map', level=2)
add_bullet('**Filter by category:** Tap a pick name in the filter bar at the top to show only that category. Tap "All" to show everything.')
add_bullet('**Search a new area:** Pan or zoom the map, then tap the **"Search This Area"** button that appears at the bottom. The app will fetch new ghost pin suggestions and run a quick pre-screen of their homepages.')
add_bullet('**Center on your location:** Tap the location button (top right) to snap back to where you are.')
add_bullet('**Refresh spots:** Tap the refresh button (top right) to reload confirmed spots from the database.')

doc.add_heading('Tapping a Ghost Pin', level=2)
add_body('When you tap any ghost pin, the map centers it and a detail sheet slides up. Here\'s what happens:')
add_bullet('**Pass 1:** The app downloads the venue\'s homepage and menu subpages, scanning for your pick\'s keywords.')
add_bullet('**Pass 2:** If Pass 1 didn\'t find a match, it searches the web for your keywords on that venue\'s domain.')
add_bullet('**Pass 3:** (Only for venues with no website) Searches the broader web for your pick at that venue name.')
add_body('While the primary pick is being checked, all your other picks are scanned against the same cached web pages at no extra cost.')

add_body('The sheet shows:')
add_bullet('**Category chips** indicating what was found (colored checkmark), what had a partial match (amber question mark), and what wasn\'t found (gray)')
add_bullet('A **result banner** explaining what the web check found')
add_bullet('A **disclaimer** reminding you that web results aren\'t perfect')

add_body('From this sheet you can:')
add_bullet('**"Yes, add it to Flezcal!"** \u2014 if you know this place has what you\'re looking for')
add_bullet('**"View in Apple Maps"** \u2014 open the venue in Apple Maps for directions')
add_bullet('**"Not accurate, remove pin"** \u2014 dismiss the ghost pin if it\'s not relevant')

add_divider()

# ============================================================
# 6. ADDING A SPOT
# ============================================================

doc.add_heading('6. Adding a Spot', level=1)

add_body('There are two ways to add a spot to Flezcal:')

doc.add_heading('From a Ghost Pin', level=2)
add_body('Tap a ghost pin, then tap **"Yes, add it to Flezcal!"** This pre-fills the venue information and any categories the web check confirmed.')

doc.add_heading('From Explore Search', level=2)
add_bullet('Go to the **Spots** tab and switch to **Explore** mode')
add_bullet('Search for a venue by name')
add_bullet('Tap any search result to open the detail sheet')
add_bullet('Tap **"Yes, add it to Flezcal!"** to add it')

doc.add_heading('The Confirm Spot Screen', level=2)
add_body('After selecting a venue, the Confirm Spot screen shows:')
add_bullet('A map preview with the venue\'s location')
add_bullet('The venue name and address, verified via Apple Maps')
add_bullet('If the spot already exists on Flezcal, a banner tells you \u2014 your additions will merge with the existing entry')

doc.add_heading('Adding Offerings', level=2)
add_body('Each category has a specific type of offering you can list. These help other users know exactly what\'s available:')

add_table(
    ['Category', 'Offering Type', 'Examples'],
    [
        ['Mezcal', 'Mezcal Brands', 'Del Maguey, Vago, Bozal'],
        ['Flan', 'Flan Styles', 'Classic, Coconut, Cheese Flan'],
        ['Bourbon', 'Bourbon Brands', 'Maker\'s Mark, Woodford Reserve'],
        ['Oysters', 'Oyster Varieties', 'Wellfleet, Kumamoto, Blue Point'],
        ['Wood-Fired Pizza', 'Pizza Styles', 'Margherita, Marinara, Diavola'],
        ['Handmade Tortillas', 'Tortilla Types', 'Corn, Flour, Blue Corn'],
        ['Pho', 'Pho Types', 'Tai (Rare Beef), Dac Biet (Special)'],
        ['Pozole', 'Pozole Styles', 'Rojo, Verde, Blanco'],
        ['Lobster Rolls', 'Roll Styles', 'Maine Style, Connecticut Style'],
        ['Caviar', 'Caviar Types', 'Osetra, Beluga, Paddlefish'],
    ]
)

add_body('Enter one item per field. Tap **"Add Another"** to list more. Offerings are optional but helpful.')
add_body('Tap **"Add This Spot"** (or **"Update This Spot"** if it already exists) to save.')

add_divider()

# ============================================================
# 7. SPOT DETAILS
# ============================================================

doc.add_heading('7. Spot Details', level=1)

add_body('Tap any confirmed spot on the map or in the Spots list to see its full detail page.')

doc.add_heading('What You\'ll See', level=2)
add_bullet('**Photo:** A map snapshot showing the venue\'s location.')
add_bullet('**Category badges:** Colored chips showing all categories this spot is tagged with.')
add_bullet('**Name and address**')
add_bullet('**Average rating:** Displayed as a flan emoji bar (instead of stars) with a numeric average and the number of ratings.')
add_bullet('**Fun badges:** "Hidden Gem" (for highly-rated spots with few ratings) or "New" (added within the last 30 days).')

doc.add_heading('Verification Section', level=2)
add_body('Under "Help verify this spot," each category has thumbs-up and thumbs-down buttons. If you\'ve been to this place and can confirm whether they serve a particular item, vote! A confirmation percentage shows the community consensus (e.g., "85% of users confirm Mezcal").')

doc.add_heading('Offerings', level=2)
add_body('Listed by category. For example, a mezcal spot might show brands like "Del Maguey," "Vago," and "Bozal." You can add more offerings by tapping the + button if the category isn\'t locked by the owner.')

doc.add_heading('Adding Another Category', level=2)
add_body('If you know a spot serves something that isn\'t listed yet, tap **"Add Another Category"** at the bottom. This runs a website check for the new category and lets you add it to the spot.')

doc.add_heading('Actions', level=2)
add_bullet('**Open in Maps:** Get directions via Apple Maps')
add_bullet('**Rate:** Give your rating')
add_bullet('**Report Spot:** Flag inaccurate or inappropriate listings')
add_bullet('**Report as Permanently Closed:** Let the community know if a place has shut down')

add_divider()

# ============================================================
# 8. RATING A SPOT
# ============================================================

doc.add_heading('8. Rating a Spot', level=1)

add_body('Flezcal uses a flan-based rating scale instead of stars:')

add_table(
    ['Rating', 'Label', 'Meaning'],
    [
        ['\U0001f36e', 'Meh', 'Wouldn\'t order it again'],
        ['\U0001f36e\U0001f36e', 'Decent', 'Fine, nothing special'],
        ['\U0001f36e\U0001f36e\U0001f36e', 'Legit', 'Solid, the real deal'],
        ['\U0001f36e\U0001f36e\U0001f36e\U0001f36e', 'Fire', 'One of the best I\'ve had'],
        ['\U0001f36e\U0001f36e\U0001f36e\U0001f36e\U0001f36e', 'Obsessed', 'Best I\'ve ever had, I\'m obsessed'],
    ]
)

add_body('**To rate a spot:**')
add_bullet('Open the spot detail page')
add_bullet('Tap the **"Rate"** button')
add_bullet('If the spot has multiple categories, select which one you\'re rating')
add_bullet('Choose your rating level')
add_bullet('Submit')
add_body('Only rate items you\'ve personally tried at that location.')

add_divider()

# ============================================================
# 9. THE SPOTS TAB
# ============================================================

doc.add_heading('9. The Spots Tab', level=1)

add_body('The **Spots** tab (list icon) has two modes, toggled by a segmented control at the top:')

doc.add_heading('Community Mode', level=2)
add_body('Shows all confirmed Flezcal spots from the database. Use the filter pills to narrow by category. Each row shows:')
add_bullet('Category icons')
add_bullet('Spot name')
add_bullet('Status (rating bar, verification status, or "New")')
add_bullet('Distance from your location in miles')

doc.add_heading('Explore Mode', level=2)
add_body('A live search powered by Apple Maps. Type a venue name to find any restaurant, bar, or store \u2014 even ones not yet on Flezcal.')
add_body('Results show:')
add_bullet('An icon indicating the venue type (fork/knife for restaurants, wineglass for bars, mug for breweries, etc.)')
add_bullet('Green highlight if the homepage pre-screen found a keyword match')
add_bullet('Venue name, address, and distance')
add_body('Tap any result to open the ghost pin detail sheet with an automated website check. From there you can add it to Flezcal or tap **"Show on Map"** to see it on the map.')
add_body('**Searching other cities:** You can search for venues in any location, not just near you.')

add_divider()

# ============================================================
# 10. LEADERBOARD
# ============================================================

doc.add_heading('10. Leaderboard', level=1)

add_body('The **Leaderboard** tab (trophy icon) tracks community contributions.')

doc.add_heading('How Scoring Works', level=2)

add_table(
    ['Action', 'Points'],
    [
        ['Add a new spot', '+10'],
        ['Rate a spot', '+5'],
        ['Category identified on a spot', '+3'],
        ['Log an offering (brand, style, etc.)', '+1'],
        ['Verify a spot', '+1'],
    ]
)

doc.add_heading('Contributor Ranks', level=2)
add_body('As you earn points, you climb through the ranks:')

add_table(
    ['Rank', 'Points Needed'],
    [
        ['Turista', '0'],
        ['Chilango', '1'],
        ['Flan\u00e1tico', '20'],
        ['Mezcalero', '50'],
        ['Conocedor', '100'],
        ['Leyenda CDMX', '200'],
        ['Inmortal', '500'],
    ]
)

add_body('**Brand Collector badge:** Log 10+ mezcal brands to earn this special badge.')

add_divider()

# ============================================================
# 11. PROFILE
# ============================================================

doc.add_heading('11. Profile', level=1)

add_body('The **Profile** tab (person icon) shows your account info and contributions.')
add_body('When signed in, you\'ll see:')
add_bullet('Your display name and email')
add_bullet('Your contributor rank and points')
add_bullet('Contribution counts (spots added, flan spots, mezcal spots, brands listed, ratings given)')
add_bullet('**Change Display Name** \u2014 this is the name that appears on the leaderboard')
add_bullet('**What\'s New** \u2014 revisit the welcome/feature screen')
add_bullet('**Version** \u2014 current app version')
add_bullet('**Privacy Policy** \u2014 link to flezcal.app/privacy')
add_bullet('**Sign Out** and **Delete Account** options')

add_divider()

# ============================================================
# 12. FOR VENUE OWNERS
# ============================================================

doc.add_heading('12. For Venue Owners', level=1)

add_body('If you own or manage a spot that\'s on Flezcal, you can update your listing just like any other user \u2014 for free. To go further, contact support@flezcal.app to become an **Owner Verified** spot. Owner verification gives you:')
add_bullet('A verified badge on your listing')
add_bullet('The ability to lock your menu details so others can\'t modify them')
add_bullet('A section to list your featured brands')
add_bullet('A reservation link on your spot page')

add_divider()

# ============================================================
# 13. TIPS AND TRICKS
# ============================================================

doc.add_heading('13. Tips and Tricks', level=1)

add_bullet('**Shake your phone** for a surprise.')
add_bullet('**Green ghost pins are your best bet** \u2014 the app already found keyword matches on their website. Tap those first.')
add_bullet('**Edit your search terms** if you\'re not getting good results. Adding terms that restaurants actually use on their menus makes the web check more accurate.')
add_bullet('**The web check scans menu subpages too**, not just the homepage. It follows links to pages like "/menu" or "/food" to find your items.')
add_bullet('**Your picks shape everything.** Change your picks and the map, ghost pins, and filters all update to match.')
add_bullet('**You can search other cities** in the Explore tab \u2014 useful for trip planning.')
add_bullet('**Rate only what you\'ve tried in person.** The rating scale is about the specific item at that specific place, not the restaurant overall.')

add_divider()

# ============================================================
# 14. CATEGORY REFERENCE
# ============================================================

doc.add_heading('14. Category Reference', level=1)

add_body('The complete list of all 23 built-in categories, their emojis, and the default search terms used for website scanning.')

doc.add_heading('Launch Trio', level=2)

add_table(
    ['Category', 'Emoji', 'Default Search Terms'],
    [
        ['Mezcal', '\U0001f943', 'mezcal, mezcaleria, mezcales, mezcal list, mezcal menu, mezcal selection, agave spirits'],
        ['Flan', '\U0001f36e', 'flan, flan casero'],
        ['Handmade Tortillas', '\U0001fad3', 'handmade tortillas, tortillas hechas a mano, fresh tortillas, tortilleria, house-made tortillas, homemade tortillas'],
    ]
)

doc.add_heading('Drinks', level=2)

add_table(
    ['Category', 'Emoji', 'Default Search Terms'],
    [
        ['Bourbon', '\U0001f943', 'bourbon, kentucky bourbon, small batch bourbon, single barrel bourbon, bourbon selection'],
        ['Fernet Branca', '\U0001f33f', 'fernet, fernet branca, fernet-branca, amaro, digestif'],
        ['New England IPA', '\U0001f37a', 'new england ipa, neipa, hazy ipa, juicy ipa, hazy pale ale'],
        ['Single Malt Scotch', '\U0001f943', 'single malt, scotch whisky, single malt scotch, speyside, islay, highland scotch'],
    ]
)

doc.add_heading('Savory', level=2)

add_table(
    ['Category', 'Emoji', 'Default Search Terms'],
    [
        ['Peameal Bacon', '\U0001f953', 'peameal bacon, peameal, canadian bacon, back bacon, cornmeal bacon'],
        ['Wood-Fired Pizza', '\U0001f355', 'wood fired, wood-fired, wood oven, brick oven, neapolitan, napoletana'],
        ['Paella', '\U0001f958', 'paella, paella valenciana, paella mixta, arroz, bomba rice'],
        ['Oysters', '\U0001f9aa', 'oysters, oyster bar, fresh oysters, oyster selection, shucked'],
        ['Pho', '\U0001f372', 'pho, pho menu'],
        ['Pozole', '\U0001f372', 'pozole, pozole rojo, pozole verde, pozole blanco, pozoleria'],
        ['Tartare', '\U0001f969', 'tartare, steak tartare, beef tartare, tuna tartare, salmon tartare'],
        ['Fugu', '\U0001f421', 'fugu, pufferfish, blowfish, fugu sashimi, tessa'],
        ['Bibimbap', '\U0001f35a', 'bibimbap, dolsot bibimbap, stone pot bibimbap, mixed rice bowl'],
        ['Iberico Ham', '\U0001f356', 'iberico, jamon iberico, pata negra, bellota'],
        ['Caviar', '\U0001fae7', 'caviar, osetra, beluga caviar, sturgeon caviar, caviar service'],
        ['Pierogi', '\U0001f95f', 'pierogi, pierog, pierogy, pierogies, polish dumplings, ruskie'],
        ['Lobster Rolls', '\U0001f99e', 'lobster roll, lobster rolls, lobster sandwich, maine lobster roll, connecticut lobster roll'],
        ['Smashburgers', '\U0001f354', 'smashburger, smash burger, smashed burger, smash patty, crispy edges'],
    ]
)

doc.add_heading('Sweets & Specialty', level=2)

add_table(
    ['Category', 'Emoji', 'Default Search Terms'],
    [
        ['Maple Syrup', '\U0001f341', 'maple syrup, pure maple, maple sugar, sugar shack, grade a maple'],
        ['Artisan Chocolate', '\U0001f36b', 'artisan chocolate, bean to bar, craft chocolate, single origin chocolate, chocolatier, cacao'],
    ]
)

# ============================================================
# BACK COVER
# ============================================================

doc.add_page_break()

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Flezcal')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0xE8, 0x6B, 0x1A)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Life is short. Eat flan.')
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('support@flezcal.app')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xE8, 0x6B, 0x1A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('flezcal.app')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xE8, 0x6B, 0x1A)

# ============================================================
# SAVE
# ============================================================

output_path = '/Users/peterwojciechowski/Library/Mobile Documents/com~apple~CloudDocs/Claude coding/Flezcal/Flezcal_User_Manual.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
